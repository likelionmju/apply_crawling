from selenium.webdriver import Chrome, ChromeOptions
from bs4 import BeautifulSoup
from pathlib import Path
from requests import get as request_get
from zipfile import ZipFile
from shutil import copyfileobj
from docx import Document
from applicant import Applicant
from email.mime.text import MIMEText
import dill
import smtplib
import main
import pandas

domain = "https://apply.likelion.org"
apply_url = f"{domain}/apply"
univ_url = f"{apply_url}/univ"
applicant_url = f"{apply_url}/applicant"

id_path = "id_username"
password_selector = "id_password"
login_selector = "//button[@type='submit']"
applicant_info_container_selector = "#likelion_num"
applicant_answer_container_selector = ".answer_view > .applicant_detail_page"
applicants_selector = f"{applicant_info_container_selector} > div.applicant_page > a"
applicant_name_selector = "h3"
applicant_info_selector = "div.row"
applicant_answer_selector = "div.m_mt"
link_attr = "href"

html_parser = "html.parser"

sns_list = ("facebook", "instagram", "twitter")
img_extensions = (".png", ".jpg", ".jpeg", ".PNG", ".JPG", ".JPEG")
doc_extensions = (".pdf", ".docx", ".hwp")
archive_extension = ".zip"


# Important! This source is use ChromeDriver please download and move to this file's directory

# If link is sns(facebook, instagram, twitter) return True
def is_sns(link) -> bool:
    for lk in sns_list:
        if lk in link:
            return True
    return False


# If file is image(png, jpg, jpeg) return True
def is_img(img: Path) -> bool:
    return img.suffix in img_extensions


# If file is document(docx, pdf, hwp) return True
def is_doc(doc: Path) -> bool:
    return doc.suffix in doc_extensions


# If file is archived return True
def is_archive(archive: Path) -> bool:
    return archive.suffix == archive_extension


# Unzip file to (param: to)
def unzip(target: Path, to) -> None:
    with ZipFile(target) as zip_file:
        info = zip_file.infolist()
        for file in info:
            t = Path(file.filename)
            if t.is_dir():
                continue
            if not (is_doc(t) or is_img(t)):
                # If file is img or document continue
                continue
            if file.flag_bits != 2048:
                # If not utf-8
                # This code is fix Korean file name error
                file.filename = file.filename.encode("cp437").decode("cp949")
            zip_file.extract(file, to)


# File rename if file is image or document
def reformat_file(file: Path) -> None:
    if is_img(file):
        file.rename(f"{file.parent}/시간표{file.suffix}")
    elif is_doc(file):
        file.rename(f"{file.parent}/포트폴리오{file.suffix}")


# Download file by url
def download_file_by_url(u, save_path: Path) -> None:
    # source: https://stackoverflow.com/a/9419208
    r = request_get(u, stream=True)
    with open(save_path, 'wb') as fd:
        r.raw.decode_content = True
        copyfileobj(r.raw, fd)


# Download applicant's extra files
def download_applicant_file(applicant: Applicant) -> None:
    if not applicant.root_dir.exists():
        applicant.root_dir.mkdir()
    if not applicant.has_file():
        return
    target_file = Path(f"{applicant.root_dir}/{applicant.cdn_file.split('/')[-1]}")
    download_file_by_url(applicant.cdn_file, target_file)
    if is_archive(target_file):
        archive_dir = Path(f"{applicant.root_dir}/시간표 및 포트폴리오")
        unzip(target_file, archive_dir)
        target_file.unlink()
    else:
        reformat_file(target_file)


# Export applicant's data
def export_docx(applicant: Applicant) -> None:
    docx = Document()
    if applicant.is_exclude:
        return
    for info in applicant.information_stringify():
        docx.add_paragraph(info)
    for idx, qna in enumerate(zip(main.secrets["QUESTIONS"], applicant.answers), 1):
        docx.add_paragraph("")
        docx.add_paragraph(f"{idx}. {qna[0]}")
        docx.add_paragraph(qna[1])
    docx.save(f"{applicant.root_dir}/지원서.docx")


def login(admin_id: str, admin_password: str, with_headless: bool = True) -> dict:
    driver_options = ChromeOptions()
    driver_options.headless = with_headless
    with Chrome(executable_path="../chromedriver", options=driver_options) as driver:
        driver.get(apply_url)
        driver.find_element_by_id(id_path).send_keys(admin_id)
        driver.find_element_by_id(password_selector).send_keys(admin_password)
        driver.find_element_by_xpath(login_selector).submit()
        return driver.get_cookies()


# Get page source at https://apply.likelion.org/apply/univ/~
def request_univ_page_source(univ_code: str, login_info: dict) -> str:
    return request_get(f"{univ_url}/{univ_code}", cookies=login_info).text


# Get all applicant primary keys(maybe)
def extract_all_applicant_pks(univ_page_source: str) -> list:
    univ_page = BeautifulSoup(univ_page_source, features=html_parser)
    return [applicant.get(link_attr).split("/")[-1] for applicant in univ_page.select(applicants_selector)]


# Get applicant's source page at https://apply.likelion.org/apply/applicant/~
def request_applicant_source(applicant_pk: str, login_info: dict) -> str:
    res = request_get(f"{applicant_url}/{applicant_pk}", cookies=login_info)
    return res.text


# Parse applicant's source page using BeautifulSoup
def parse_applicant_page(page: str, q_count: int) -> Applicant:
    soup = BeautifulSoup(page, features=html_parser)
    applicant_info_container = soup.select_one(applicant_info_container_selector)
    applicant_answer_container = soup.select_one(applicant_answer_container_selector)

    applicant_name = applicant_info_container.find(applicant_name_selector).string
    if applicant_name in main.secrets["EXCLUDES"]:
        return Applicant.get_exclude_applicant()
    applicant_info_list = applicant_info_container.select(applicant_info_selector)
    applicant_answer_list = applicant_answer_container.select(applicant_answer_selector)
    additional = [user_info.contents[1].get(link_attr) for user_info in applicant_info_list[2:]]
    applicant_git = applicant_sns = applicant_file = None

    for item in additional:
        if item is None:
            continue
        if "git" in item:
            applicant_git = item
        elif is_sns(item):
            applicant_sns = item
        elif "cdn" in item:
            applicant_file = item

    applicant = Applicant(name=applicant_name,
                          entrance_year=applicant_info_list[0].contents[1].text,
                          major=applicant_info_list[0].contents[-2].text,
                          phone_num=applicant_info_list[1].contents[1].text,
                          email=applicant_info_list[1].contents[-2].text,
                          answers=[applicant_answer_list[idx].contents[1].text for idx in
                                   range(q_count)],
                          git=applicant_git,
                          sns=applicant_sns,
                          cdn_file=applicant_file,
                          )
    applicant.format_phone_num()

    return applicant


def pickle_applicant(applicant: Applicant) -> None:
    if applicant.is_exclude:
        return
    pkl_dir = Path(f"../applicant/{applicant.name}.pkl")
    with open(pkl_dir, "wb") as pkl:
        dill.dump(applicant, pkl)


def unpickle_all_applicant():
    rv = []
    for pkl in Path("../applicant").iterdir():
        if pkl.suffix == ".pkl":
            with open(pkl, "rb") as p:
                rv.append(dill.loads(p.read()))
    return rv


def unpickle_applicant(name: str):
    pkl_dir = Path(f"../applicant/{name}.pkl")
    with open(pkl_dir, "rb") as pkl:
        return dill.loads(pkl.read())


def send_email_to_applicant(applicant: Applicant, is_markdown: bool):
    m_type = "html" if is_markdown else "plain"
    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.ehlo()
        smtp.login(main.secrets["EMAIL"], main.secrets["EMAIL_PASSWORD"])

        if applicant.is_pass:
            msg = MIMEText(main.data["pass_text"], m_type)
            msg["Subject"] = main.data["pass_subject"]
        else:
            msg = MIMEText(main.data["fail_text"], m_type)
            msg["Subject"] = main.data["fail_subject"]
        msg["From"] = main.secrets["EMAIL_FROM"]
        msg["To"] = applicant.email
        smtp.sendmail(main.secrets["EMAIL"], applicant.email, msg.as_string())


def gathering_applicant_data(applicants: list):
    records = []
    for applicant in applicants:
        if not applicant.is_exclude:
            records.append(applicant.information_to_dict())
    df = pandas.DataFrame.from_records(records, index=range(1, len(records) + 1))
    df = df.sort_values(["전공"])
    df.to_excel("../applicant/지원자 정보.xlsx")

